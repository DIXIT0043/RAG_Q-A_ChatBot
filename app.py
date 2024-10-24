import streamlit as st
import os
import PyPDF2
from docx import Document
from langchain_groq import ChatGroq
from langchain_huggingface import HuggingFaceEmbeddings 
from langchain_community.embeddings import OllamaEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document as LangchainDocument
import openai
import time

# Load environment variables
from dotenv import load_dotenv
load_dotenv()
os.environ['OPENAI_API_KEY'] = os.getenv("OPENAI_API_KEY")
os.environ['GROQ_API_KEY'] = os.getenv("GROQ_API_KEY")
groq_api_key = os.getenv("GROQ_API_KEY")

# Initialize embeddings and model
embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
llm = ChatGroq(groq_api_key=groq_api_key, model_name="Llama3-8b-8192")

# Prompt template
prompt = ChatPromptTemplate.from_template(
    """
    Answer the questions based on the provided context only.
    Please provide the most accurate response based on the question
    <context>
    {context}
    </context>
    Question: {input}
    """
)

def read_pdf(file):
    try:
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading PDF file: {str(e)}")
        return ""

def read_docx(file):
    try:
        doc = Document(file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading DOCX file: {str(e)}")
        return ""

def create_vector_embedding(texts):
    try:
        # Create LangChain Document objects
        documents = [LangchainDocument(page_content=text) for text in texts if text.strip()]
        
        if not documents:
            st.warning("No valid text content found in the documents.")
            return
        
        st.session_state.text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        st.session_state.final_documents = st.session_state.text_splitter.split_documents(documents[:50])
        st.session_state.vectors = FAISS.from_documents(st.session_state.final_documents, embeddings)
        return True
    except Exception as e:
        st.error(f"Error creating vector embeddings: {str(e)}")
        return False

def main():
    st.title("RAG Document Q&A With Groq And Llama3")

    # Initialize session state if needed
    if 'processed_texts' not in st.session_state:
        st.session_state.processed_texts = []

    # Button to upload a file
    uploaded_files = st.file_uploader("Upload PDF or Document", type=["pdf", "docx"], accept_multiple_files=True)

    # Process uploaded files
    if uploaded_files:
        st.session_state.processed_texts = []
        for uploaded_file in uploaded_files:
            if uploaded_file.type == "application/pdf":
                text = read_pdf(uploaded_file)
                if text:
                    st.session_state.processed_texts.append(text)
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = read_docx(uploaded_file)
                if text:
                    st.session_state.processed_texts.append(text)
        
        if st.session_state.processed_texts:
            if create_vector_embedding(st.session_state.processed_texts):
                st.success("Documents processed and vectors created successfully.")
        else:
            st.warning("No valid documents were processed.")

    # Ask a question button
    user_prompt = st.text_input("Enter your question from the uploaded document")

    if st.button("Ask"):
        if 'vectors' in st.session_state:
            try:
                document_chain = create_stuff_documents_chain(llm, prompt)
                retriever = st.session_state.vectors.as_retriever()
                retrieval_chain = create_retrieval_chain(retriever, document_chain)

                start = time.process_time()
                response = retrieval_chain.invoke({'input': user_prompt})
                st.write(f"Response time: {time.process_time() - start}")

                st.write(response['answer'])
            except Exception as e:
                st.error(f"Error processing question: {str(e)}")
        else:
            st.warning("Please upload a document first.")

if __name__ == "__main__":
    main()